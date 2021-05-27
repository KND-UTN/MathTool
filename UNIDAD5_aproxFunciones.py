__author__ = 'Leila Spini, Ignacio Pieve, Sofia Cibello, Juan Pablo Donalisio'
__version__ = 2.0

from io import StringIO, BytesIO

import pandas as pd
import numpy as np
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from sympy import *
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
import matplotlib.pyplot as plt


def menu():
    vals = pedir_valores()
    funcs = pedir_funciones()
    constructor(vals, funcs)


def menuTest(numero):
    x = Symbol('x')
    y = Symbol('y')
    c = Symbol('c')
    e = np.exp(1)
    if numero == 1:
        valoresy = [0, 2, 3, 5, 8, 12, 12.5]
        valoresx = [0, 1, 1, 2, 3, 4.5, 5]
        vals = pd.DataFrame({"Y": valoresy, "X": valoresx})
        funcs = [eval("c*x+c"), eval("c*x**2+c*x+c"), eval("c*e**x+c"), eval("c*ln(x+1)+c*x")]
    if numero == 2:
        valoresy = [4, 2, 1.5, 3, 2.9, 2.7, 2.2, 3.8, 7.6]
        valoresx = [-3, -2.1, 1, 2, 2.5, 5.5, 6.7, 7.5, 8.1]
        vals = pd.DataFrame({"Y": valoresy, "X": valoresx})
        funcs = [eval("c*x**3+c*x"), eval("c*sin(0.2*x)+c")]
    if numero == 3:
        valoresx = [.15, .3, .5, .6]
        valoresy = [1.35, 2.3, 3.7, 2.5]
        vals = pd.DataFrame({"Y": valoresy, "X": valoresx})
        funcs = [eval("c*x**2 + c*x + c"), eval("c*x**2  + c*ln(x)"), eval("3.5*x+1.1")]
    constructor(vals, funcs)


def pedir_valores():
    cant = int(input('Cuantos valores desea ingresar?: '))
    valoresy = []
    valoresx = []
    for a in range(cant):
        valoresy.append(float(input('Ingrese el valor ' + str(a) + ' en y: ')))
    for a in range(cant):
        valoresx.append(float(input('Ingrese el valor ' + str(a) + ' en x: ')))
    m = pd.DataFrame({"Y": valoresy, "X": valoresx})
    return m


def pedir_funciones():
    x = Symbol('x')
    y = Symbol('y')
    c = Symbol('c')
    e = np.exp(1)
    cant = int(input('Cuantas funciones desea ingresar?: '))
    funciones = []
    for a in range(cant):
        f = eval(input('Ingrese la funcion ' + str(a + 1) + ': ').replace("^", "**"))
        funciones.append(f)
    return funciones


def constructor(valores, ecuaciones):
    x = Symbol('x')
    y = Symbol('y')
    c = Symbol('c')
    e = np.exp(1)
    ecuaciones_resueltas = []
    matrices_sin_datos = []
    matrices_con_datos = []
    soluciones = []
    terminos = []
    desviaciones = []
    for a in range(len(ecuaciones)):
        argumentos = (ecuaciones[a]).args
        cant_filas = len(argumentos)
        matriz = [[0 for x in range(cant_filas + 1)] for y in range(cant_filas)]
        matriz2 = [[0 for x in range(cant_filas + 1)] for y in range(cant_filas)]
        for i in range(cant_filas):
            terminos.append('C' + str(cant_filas - i))
            for j in range(cant_filas + 1):
                if j == cant_filas:
                    matriz[i][j] = ((argumentos[i] / c) * y)
                else:
                    matriz[i][j] = ((argumentos[i] / c) * (argumentos[j] / c))

                total_esta_celda = 0
                for k in range(len(valores["X"])):
                    total_esta_celda += matriz[i][j].subs(x, valores['X'][k]).subs(y, valores['Y'][k])
                matriz2[i][j] = total_esta_celda

        matrices_sin_datos.append(matriz)
        matrices_con_datos.append(matriz2)

        solucionar = Matrix(matriz2)
        solucion = solve_linear_system_LU(solucionar, terminos)
        soluciones.append(solucion)
        terminos = []
        ecuacion_resuelta = []
        cont = 0
        for i in range(len(argumentos)):
            if c in argumentos[i].free_symbols:
                ecuacion_resuelta.append(argumentos[i].subs(c, solucion['C' + str(cant_filas - cont)]))
                cont += 1
            else:
                ecuacion_resuelta.append(argumentos[i])
        unida = 0
        for i in range(len(ecuacion_resuelta)):
            unida += ecuacion_resuelta[i]
        ecuaciones_resueltas.append(unida)

        desviacion = 0
        for k in range(len(valores["X"])):
            t1 = ecuaciones_resueltas[a].subs(x, valores['X'][k]).subs(y, valores['Y'][k])
            t2 = valores['Y'][k] - t1
            desviacion += t2 ** 2
        desviaciones.append(desviacion)

    #mostrador(ecuaciones, matrices_sin_datos, matrices_con_datos, soluciones, ecuaciones_resueltas, valores, desviaciones)
    generadorDocumento(ecuaciones, matrices_sin_datos, matrices_con_datos, soluciones, ecuaciones_resueltas, valores,
                       desviaciones)


def mostrador(ecs, matricessd, matricescd, soluciones, ecsresueltas, datos, desv):
    print('-' * 75)
    print('Datos:')
    print(datos)
    print('')
    print('Ecuaciones planteadas:')
    for a in ecs:
        print('F(x) = ' + str(a).replace("**", "^"))

    c = 0
    for a in ecs:
        print('\n' + '-' * 75)
        c += 1
        print('Ecuacion ' + str(c) + ': ' + str(a))
        print('\nMatriz con datos (Sin resolver):')
        temp = Matrix(matricessd[c - 1])
        temp = np.array(temp.transpose())
        msr = {}
        for b in range(len(matricessd[c - 1][0])):
            if b == len(matricessd[c - 1][0]) - 1:
                msr['Y'] = temp[b]
            else:
                msr['C' + str(b + 1)] = temp[b]
        temp = pd.DataFrame(msr)
        print(temp)

        print('\nMatriz con datos (Resuelta):')
        temp = Matrix(matricescd[c - 1])
        temp = np.array(temp.transpose())
        msr = {}
        for b in range(len(matricescd[c - 1][0])):
            if b == len(matricescd[c - 1][0]) - 1:
                msr['Y'] = temp[b]
            else:
                msr['C' + str(b + 1)] = temp[b]
        temp = pd.DataFrame(msr)
        print(temp)

        print('\nConjunto de soluciones: ')
        for i in soluciones[c - 1]:
            print(str(i) + ': ' + str(soluciones[c - 1][i]))
        print('\nEcuacion final: ' + str(ecsresueltas[c - 1]))
        print('Desviacion: ' + str(desv[c - 1]))


def generadorDocumento(ecs, matricessd, matricescd, soluciones, ecsresueltas, datos, desv):

    document = Document()
    document.add_heading('Aproximación de Funciones', level=0)
    document.add_heading('Datos de Entrada', level=1)
    document.add_paragraph()


    # ----- ARMADO DE TABLA -----
    anchoCeldas = Inches(.75)
    table = document.add_table(rows=1, cols=2, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    sombreado1 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    sombreado2 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    hdr_cells[0].text = 'X'
    hdr_cells[1].text = 'Y'
    hdr_cells[0].width = anchoCeldas
    hdr_cells[1].width = anchoCeldas
    hdr_cells[0]._tc.get_or_add_tcPr().append(sombreado1)
    hdr_cells[1]._tc.get_or_add_tcPr().append(sombreado2)
    hdr_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for a in range(len(datos)):
        row_cells = table.add_row().cells
        row_cells[0].text = str(datos['X'][a])
        row_cells[1].text = str(datos['Y'][a])
        row_cells[0].width = anchoCeldas
        row_cells[1].width = anchoCeldas
        row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # ----- MUESTRA DE FUNCIONES INGRESADAS POR EL USUARIO -----
    document.add_paragraph()
    p = document.add_paragraph()
    p.add_run("Ecuaciones planteadas:").bold = True
    for a in range(len(ecs)):
        p = document.add_paragraph()
        p.add_run('F' + str(a) + '(x) = ').bold = True
        p.add_run(str(ecs[a]).replace("**", "^")).italic = True

    # ----- MUESTRA DE RESULTADOS -----
    c = 0
    for a in ecs:
        c += 1
        document.add_page_break()
        document.add_heading('Función ' + str(c), level=0)
        p = document.add_paragraph()
        p.add_run('F' + str(c) + '(x) = ').bold = True
        p.add_run(str(a).replace("**", "^")).italic = True

        document.add_paragraph()
        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.add_run("Matriz sin resolver").bold = True

        msr = matrizInvertidaDiagonalConCs(matricessd[c - 1])
        temp = pd.DataFrame(msr)

        anchoCeldas = Inches(.75)
        table = document.add_table(rows=0, cols=len(temp.columns), style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for fila in range(len(temp)):
            row_cells = table.add_row().cells
            for columna in range(len(temp.columns)):
                texto = str(temp[temp.columns[columna]][fila]).replace("**", "^")
                if columna == len(temp.columns) - 2:
                    sombreado = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
                    row_cells[columna]._tc.get_or_add_tcPr().append(sombreado)
                else:
                    if texto == "1":
                        texto = "n"
                    else:
                        texto = "Σ" + texto

                row_cells[columna].text = texto
                row_cells[columna].width = anchoCeldas


        document.add_paragraph()
        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.add_run("Matriz resuelta").bold = True

        msr = matrizInvertidaDiagonalConCs(matricescd[c - 1])
        temp = pd.DataFrame(msr)

        anchoCeldas = Inches(.75)
        table = document.add_table(rows=0, cols=len(temp.columns), style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for fila in range(len(temp)):
            row_cells = table.add_row().cells
            for columna in range(len(temp.columns)):
                row_cells[columna].text = str(temp[temp.columns[columna]][fila])
                row_cells[columna].width = anchoCeldas
                if columna == len(temp.columns) - 2:
                    sombreado = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
                    row_cells[columna]._tc.get_or_add_tcPr().append(sombreado)

        document.add_paragraph()
        document.add_paragraph()
        p = document.add_paragraph()
        p.add_run("Conjunto de Soluciones: ").bold = True
        for i in reversed(soluciones[c - 1]):
            p = document.add_paragraph()
            p.add_run(str(i) + ': ').bold = True
            p.add_run(str(soluciones[c - 1][i]))

        document.add_paragraph()
        p = document.add_paragraph()
        p.add_run("Ecuacion final: ").bold = True
        p.add_run(str(ecsresueltas[c - 1]).replace("**", "^"))
        p = document.add_paragraph()
        p.add_run("Desviacion: ").bold = True
        p.add_run(str(desv[c - 1]))

    document.add_page_break()
    document.add_heading("Gráficos", level=0)
    # Grafico
    plt.scatter(datos['X'], datos['Y'], alpha=0.5)
    x = np.arange(datos['X'].min(), datos['X'].max(), (datos['X'].max() - datos['X'].min()) / 100)
    numeroEcuacion = 0
    leyendas = []
    for ecuacion in ecsresueltas:
        numeroEcuacion += 1
        y = eval(str(ecuacion).replace("sin", "msin").replace("cos", "mcos").replace("log", "mln"))
        plt.plot(x, y, label=str(ecuacion))
        leyendas.append("Ecuación " + str(numeroEcuacion))
    plt.legend(leyendas)
    diagramaDispersion = BytesIO()
    plt.savefig(diagramaDispersion)
    document.add_picture(diagramaDispersion, width=Inches(6))
    plt.clf()
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    numeroEcuacion = 0
    for ecuacion in ecsresueltas:
        plt.scatter(datos['X'], datos['Y'], alpha=0.5)
        numeroEcuacion += 1
        y = eval(str(ecuacion).replace("sin", "msin").replace("cos", "mcos").replace("log", "mln"))
        plt.plot(x, y, label=str(ecuacion))
        leyendas = ["Ecuación " + str(numeroEcuacion)]
        plt.legend(leyendas)
        diagramaDispersion = BytesIO()
        plt.savefig(diagramaDispersion)
        document.add_picture(diagramaDispersion, width=Inches(5))
        plt.clf()
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_page_break()
    document.add_heading("Conclusión", level=0)
    p = document.add_paragraph()
    p.add_run('La función con menor desviación (óptima) es la ')
    optima = min(range(len(desv)), key=desv.__getitem__)
    p.add_run('Función ' + str(optima + 1) + ":").bold = True
    p = document.add_paragraph()
    p.add_run('F' + str(optima + 1) + "(x): ").bold = True
    p.add_run(str(ecs[optima]).replace("**", "^"))
    p = document.add_paragraph()
    p.add_run('F' + str(optima + 1) + "(x): ").bold = True
    p.add_run(str(ecsresueltas[optima]).replace("**", "^"))
    p = document.add_paragraph()
    p.add_run("Desviación: ").bold = True
    p.add_run(str(desv[optima]))

    plt.scatter(datos['X'], datos['Y'], alpha=0.5)
    y = eval(str(ecsresueltas[optima]).replace("sin", "msin").replace("cos", "mcos").replace("log", "mln"))
    plt.plot(x, y)
    leyendas = ["Ecuación " + str(optima + 1)]
    plt.legend(leyendas)
    diagramaDispersion = BytesIO()
    plt.savefig(diagramaDispersion)
    document.add_picture(diagramaDispersion, width=Inches(3.5))
    plt.clf()
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = document.add_paragraph()

    p.add_run('La función con mayor desviación (peor función) es la ')
    optima = max(range(len(desv)), key=desv.__getitem__)
    p.add_run('Función ' + str(optima + 1) + ":").bold = True
    p = document.add_paragraph()
    p.add_run('F' + str(optima + 1) + "(x): ").bold = True
    p.add_run(str(ecs[optima]).replace("**", "^"))
    p = document.add_paragraph()
    p.add_run('F' + str(optima + 1) + "(x): ").bold = True
    p.add_run(str(ecsresueltas[optima]).replace("**", "^"))
    p = document.add_paragraph()
    p.add_run("Desviación: ").bold = True
    p.add_run(str(desv[optima]))

    plt.scatter(datos['X'], datos['Y'], alpha=0.5)
    y = eval(str(ecsresueltas[optima]).replace("sin", "msin").replace("cos", "mcos").replace("log", "mln"))
    plt.plot(x, y)
    leyendas = ["Función " + str(optima + 1)]
    plt.legend(leyendas)
    diagramaDispersion = BytesIO()
    plt.savefig(diagramaDispersion)
    document.add_picture(diagramaDispersion, width=Inches(3.5))
    plt.clf()
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.save('aproximacion.docx')


def matrizInvertidaDiagonalConCs(entrada):
    salida = []

    ultimaColumna = []
    penultimaColumna = []

    for fila in range(len(entrada)):
        ultimaColumna.insert(0, entrada[fila][len(entrada)])
        penultimaColumna.append('C' + str(fila + 1))
        entrada[fila].pop(len(entrada))

    index = 0
    for fila in reversed(entrada):
        nuevaFila = []
        for celda in reversed(fila):
            nuevaFila.append(celda)
        nuevaFila.append(penultimaColumna[index])
        nuevaFila.append(ultimaColumna[index])
        salida.append(nuevaFila)
        index += 1
    return salida


def msin(x):
    return np.sin(x)

def mcos(x):
    return np.sin(x)

def mln(x):
    return np.log(x)

if __name__ == "__main__":
    menu()
    #menuTest(1)
